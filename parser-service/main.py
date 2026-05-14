from fastapi import FastAPI, File, UploadFile
from typing import Any
from io import BytesIO
import fitz
from docx import Document
from parser_utils import changes_from_sequences, extract_paired_tracked_changes_from_docx

app = FastAPI(title="ClauseIQ Parser Service", version="0.1.0")


def _paragraphs_from_docx(raw_docx: bytes) -> list[str]:
    document = Document(BytesIO(raw_docx))
    return [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]


def _paragraphs_from_pdf(raw_pdf: bytes) -> list[str]:
    pdf = fitz.open(stream=raw_pdf, filetype="pdf")
    paragraphs: list[str] = []
    for page in pdf:
        blocks = page.get_text("blocks")
        for block in blocks:
            text = block[4].strip()
            if text:
                paragraphs.append(text)
    pdf.close()
    return paragraphs


@app.get("/health")
def health() -> dict[str, str]:
    return {"status": "ok"}


@app.post("/extract/docx-tracked")
async def extract_docx_tracked(file: UploadFile = File(...)) -> dict[str, Any]:
    data = await file.read()
    document = Document(BytesIO(data))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    tracked_changes = extract_paired_tracked_changes_from_docx(data)
    return {
        "path": "docx-tracked",
        "confidence": "high",
        "changes_detected": tracked_changes,
        "tracked_change_count": len(tracked_changes),
        "paragraph_count": len(paragraphs),
        "preview": paragraphs[:12],
    }


@app.post("/extract/docx-diff")
async def extract_docx_diff(base_file: UploadFile = File(...), redline_file: UploadFile = File(...)) -> dict[str, Any]:
    base_data = await base_file.read()
    redline_data = await redline_file.read()
    base_paragraphs = _paragraphs_from_docx(base_data)
    redline_paragraphs = _paragraphs_from_docx(redline_data)
    changes = changes_from_sequences(base_paragraphs, redline_paragraphs)
    # Clean DOCX (no tracked changes) is structurally inferred only; PRD §7 / Audit L2.
    return {
        "path": "docx-diff",
        "confidence": "low",
        "changes_detected": changes,
        "paragraph_count": len(base_paragraphs) + len(redline_paragraphs),
    }


@app.post("/extract/pdf-diff")
async def extract_pdf_diff(base_file: UploadFile = File(...), redline_file: UploadFile = File(...)) -> dict[str, Any]:
    base_data = await base_file.read()
    redline_data = await redline_file.read()
    base_paragraphs = _paragraphs_from_pdf(base_data)
    redline_paragraphs = _paragraphs_from_pdf(redline_data)
    changes = changes_from_sequences(base_paragraphs, redline_paragraphs)
    return {
        "path": "pdf-diff",
        "confidence": "low",
        "changes_detected": changes,
        "paragraph_count": len(base_paragraphs) + len(redline_paragraphs),
    }
